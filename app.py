import pandas as pd
import streamlit as st
from datetime import date, timedelta
from utils import *
from st_aggrid import GridOptionsBuilder, AgGrid, GridUpdateMode, DataReturnMode
from geopy import geocoders
from geopy.geocoders import Nominatim
import docx
import io

today = date.today()
one_week = today + timedelta(days=6)

geolocator = Nominatim(user_agent='anonymous@gmail.com', timeout=10)

st.write("""
# European commission calendar downloader
""")
date = st.sidebar.date_input("Please pick a date", value=(today, one_week))

link = f'https://commissioners.ec.europa.eu/calendar-items-commissioners_en?f%5B0%5D=ewcms_calendar_item_date%3Abt%7C' \
       f'{date[0]}T08%3A43%3A43%2B02%3A00%7C{date[1]}T08%3A43%3A43%2B02%3A00&f%5B1%5D=ewcms_calendar_status%3' \
       f'Apast&f%5B2%5D=ewcms_calendar_status%3Aupcoming'


df = get_info(link)

# st.table(data=df)
if len(df) > 0:

    df['Location'] = df['Location'].fillna('None')
    df['Commissioner'] = df['Commissioner'].fillna('None')

    gb = GridOptionsBuilder.from_dataframe(df)
    # gb.configure_pagination()
    gb.configure_selection('multiple', use_checkbox=True, groupSelectsChildren=True)
    gridOptions = gb.build()
    grid_response = AgGrid(
        df,
        gridOptions=gridOptions,
        data_return_mode=DataReturnMode.FILTERED,
        fit_columns=True,
        height=500,
        update_mode=GridUpdateMode.MODEL_CHANGED
    )

    selected = grid_response['selected_rows']
    selected = pd.DataFrame(selected)
    if len(selected) > 0:
        selected = selected[['Location']]
        selected = selected[selected['Location'] != 'None']
        if len(selected) > 0:
            selected['geocode'] = selected['Location'].apply(geolocator.geocode)
            selected = selected[-selected['geocode'].isna()]
            if len(selected) > 0:
                selected[['details', 'latlon']] = pd.DataFrame(selected['geocode'].tolist(), index=selected.index)
                selected[['lat', 'lon']] = pd.DataFrame(selected['latlon'].tolist(), index=selected.index)
                selected = selected[['Location', 'lat', 'lon']]
                st.map(selected)

else:
    st.write("""Please select a valid timeframe""")

if len(df) > 0:
    doc = docx.Document()
    for index, row in df.iterrows():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 0
        if row['Commissioner'] != 'None':
            if row['Location'] != 'None':
                run = p.add_run(f"{row['Day']}" + '/' + f"{row['Month']}" + '/' + f"{row['Year']}" + ' - ' +
                                f"{row['Commissioner'].strip()}" + ' - ' + f"{row['Location'].strip()}")
            else:
                run = p.add_run(f"{row['Day']}" + '/' + f"{row['Month']}" + '/' + f"{row['Year']}" + ' - ' +
                                f"{row['Commissioner'].strip()}")
        elif row['Location'] != 'None':
            run = p.add_run(f"{row['Day']}" + '/' + f"{row['Month']}" + '/' + f"{row['Year']}" + ' - ' +
                            f"{row['Location'].strip()}")
        else:
            run = p.add_run(f"{row['Day']}" + '/' + f"{row['Month']}" + '/' + f"{row['Year']}")

        run.bold = True
        run.font.name = 'Arial'

        run = p.add_run(f"\n")
        run.font.name = 'Arial'

        run = p.add_run(f"{row['Text']}")
        run.font.name = 'Arial'

        run = p.add_run(f"\n")
        run.font.name = 'Arial'
        run = p.add_run(f"\n")
        run.font.name = 'Arial'

    doc.save(f"CommissionCalendar{date[0]}_{date[1]}.docx")

    doc = docx.Document(f"CommissionCalendar{date[0]}_{date[1]}.docx")
    bio = io.BytesIO()
    doc.save(bio)
    if doc:
        st.sidebar.download_button('Download word file',
                                   data=bio.getvalue(),
                                   file_name=f"CommissionCalendar{date[0]}_{date[1]}.docx",
                                   mime="docx"
                                   )
